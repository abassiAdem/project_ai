from __future__ import annotations

import os
import re
import uuid
from typing import Dict, Optional

from pypdf import PdfReader
from docx import Document
from fastapi import UploadFile

from backend.app.config import settings
from backend.app.llm.groq import chat_completion


def save_upload(file: UploadFile) -> str:
    os.makedirs(settings.uploads_dir, exist_ok=True)
    ext = os.path.splitext(file.filename or "")[-1].lower()
    doc_id = f"{uuid.uuid4().hex}{ext}"
    path = os.path.join(settings.uploads_dir, doc_id)
    with open(path, "wb") as f:
        f.write(file.file.read())
    return doc_id


def _extract_text(path: str) -> str:
    if path.lower().endswith(".pdf"):
        reader = PdfReader(path)
        parts = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(parts)

    if path.lower().endswith(".docx"):
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _find_article_excerpt(text: str, query: str) -> Optional[str]:
    lowered = " ".join(text.split()).lower()
    q = query.lower()

    article_match = re.search(r"(article|المادة)\s*(\d+)", q)
    if not article_match:
        return None

    article_num = article_match.group(2)
    markers = [
        f"article {article_num}",
        f"article\s+{article_num}",
        f"المادة {article_num}",
        f"المادة\s+{article_num}",
    ]

    for marker in markers:
        pos = lowered.find(marker)
        if pos != -1:
            start = max(0, pos - 200)
            end = min(len(text), pos + 1200)
            return text[start:end].strip()
    return None


def analyze_contract(document_id: str, query: str) -> Dict[str, object]:
    path = os.path.join(settings.uploads_dir, document_id)
    if not os.path.isfile(path):
        return {"summary": "الوثيقة غير موجودة.", "notes": [], "text": "", "excerpt": ""}

    text = _extract_text(path)
    compact = " ".join(text.split())

    risk_patterns = {
        "غرامات او جزاءات": r"غرامة|جزاء|عقوبة|تعويض",
        "فسخ او انهاء": r"فسخ|انهاء|إلغاء|انقضاء",
        "تحكيم او اختصاص": r"تحكيم|اختصاص|محكمة",
        "التزامات مالية": r"مبلغ|دينار|دفعة|سداد",
        "مدة العقد": r"مدة|اجل|سنوات|اشهر",
    }

    notes = []
    for label, pattern in risk_patterns.items():
        if re.search(pattern, compact):
            notes.append(f"تم العثور على بند متعلق بـ: {label}")

    excerpt = _find_article_excerpt(text, query) or ""
    preview = compact[:4000]
    summary = "تم تحليل العقد واستخراج البنود المحتملة للمراجعة."
    return {
        "summary": summary,
        "notes": notes,
        "length": len(compact),
        "text": preview,
        "excerpt": excerpt,
    }


def _modified_dir() -> str:
    return os.path.join(settings.uploads_dir, "modified")


def modify_contract(document_id: str, instructions: str) -> Dict[str, str]:
    path = os.path.join(settings.uploads_dir, document_id)
    if not os.path.isfile(path):
        return {"status": "error", "message": "الوثيقة غير موجودة."}

    text = _extract_text(path)
    if len(text.strip()) < 50:
        return {
            "status": "error",
            "message": (
                "تعذر استخراج نص من العقد. قد يكون الملف صورة ممسوحة ضوئيا. "
                "يرجى رفع نسخة نصية أو تفعيل OCR."
            ),
        }
    prompt = (
        "انت مساعد قانوني. عدّل نص العقد حسب تعليمات المستخدم بدقة، "
        "واحتفظ بالصياغة القانونية قدر الإمكان. أعد النص كاملا بعد التعديل فقط دون شرح."
    )
    messages = [
        {"role": "system", "content": prompt},
        {
            "role": "user",
            "content": "\n\n".join(
                [
                    f"تعليمات التعديل:\n{instructions}",
                    f"نص العقد:\n{text}",
                ]
            ),
        },
    ]
    updated_text = chat_completion(messages, temperature=0.2)

    os.makedirs(_modified_dir(), exist_ok=True)
    file_id = f"modified-{uuid.uuid4().hex}.docx"
    out_path = os.path.join(_modified_dir(), file_id)

    doc = Document()
    for line in updated_text.splitlines():
        if line.strip():
            doc.add_paragraph(line)
    doc.save(out_path)

    return {"status": "ok", "file_id": file_id}
